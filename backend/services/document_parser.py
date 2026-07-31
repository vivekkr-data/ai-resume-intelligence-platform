"""Text extraction for PDF, DOCX and TXT resumes."""
from __future__ import annotations

import io
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentParseError(ValueError):
    """Raised when a resume cannot be parsed."""


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    if not file_bytes:
        raise DocumentParseError("Uploaded file is empty.")

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(
            "Unsupported file type. Upload a PDF, DOCX, or TXT file."
        )

    try:
        if extension == ".pdf":
            text = _extract_pdf(file_bytes)
        elif extension == ".docx":
            text = _extract_docx(file_bytes)
        else:
            text = file_bytes.decode("utf-8", errors="ignore")
    except Exception as exc:  # library-specific errors vary
        raise DocumentParseError(f"Could not read {filename}: {exc}") from exc

    cleaned = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if len(cleaned) < 30:
        raise DocumentParseError(
            "Very little text was extracted. The file may be scanned/image-only. "
            "Use a text-based PDF/DOCX or add OCR support."
        )
    return cleaned


def _extract_pdf(file_bytes: bytes) -> str:
    parts: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page in document:
            parts.append(page.get_text("text", sort=True))
    return "\n".join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
